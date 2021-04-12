from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)

document = Document()

# add a picture
document.add_picture(
    'myPicture.jpg', width=Inches(1.0)
    )

# name. phone number, email address
name = input('What is you name:\n')
speak('Hello' + name + 'How are you today?')
speak('Waht is your phone number?')
phone_number = input('Waht is your phone number:\n')
email = input('What is you email address:\n')


document.add_paragraph(
    name + ' |' + phone_number + ' |' + email)

# about me 
document.add_heading('About me')
about_me = input('Tell about yoursel:\n')
document.add_paragraph(about_me)

# work experience
document.add_heading('Working Experience')
p = document.add_paragraph()

company = input('Enter company name:\n')
from_date = input('From Date:\n')
to_date = input('To Date:\n')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' -- ' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ':\n')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No? \n')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company name:\n')
        from_date = input('From Date:\n')
        to_date = input('To Date:\n')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' -- ' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ':\n')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
skill = input('What is your skill? \n')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# Add more skills
while True:
    has_more_skills = input('Do you have more skills? Yes or No? \n')
    if has_more_skills.lower() == 'yes':
        skill = input('Pls enter more skills: \n')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using python programming'
                
document.save('cv.docx')